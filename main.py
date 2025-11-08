import os
from io import BytesIO
from typing import List, Optional
from fastapi import FastAPI
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, JSONResponse
from pydantic import BaseModel

# Optional DB imports (not strictly needed for this feature set)
# from database import db

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


class Skill(BaseModel):
    name: str = ""
    level: str = ""

class ExperienceItem(BaseModel):
    role: str = ""
    company: str = ""
    period: str = ""
    bullets: List[str] = []

class EducationItem(BaseModel):
    degree: str = ""
    school: str = ""
    period: str = ""
    details: str = ""

class ResumeData(BaseModel):
    name: str = ""
    title: str = ""
    email: str = ""
    phone: str = ""
    location: str = ""
    photo: Optional[str] = None
    summary: str = ""
    experience: List[ExperienceItem] = []
    education: List[EducationItem] = []
    skills: List[Skill] = []
    achievements: List[str] = []

class ExportPayload(BaseModel):
    data: ResumeData
    template: str = "clean"
    color: str = "slate"
    font: str = "inter"

class SuggestPayload(BaseModel):
    context: dict
    type: str  # 'summary' | 'bullets'


@app.get("/")
def read_root():
    return {"message": "Resume Builder Backend ready"}


@app.post("/ai/suggest")
async def ai_suggest(payload: SuggestPayload):
    # Lightweight, on-device heuristic suggestions (no external AI)
    t = payload.type
    ctx = payload.context or {}

    if t == "summary":
        title = ctx.get("title") or "professional"
        skills = ctx.get("skills") or []
        skills_str = ", ".join([s for s in skills if s])
        text = (
            f"{title.title()} with a track record of delivering measurable outcomes. "
            f"Skilled in {skills_str}. Known for clear communication, ownership, and continuous improvement. "
            f"Seeking to leverage expertise to drive impact in a high-performing team."
        ).strip()
        return {"text": text}

    if t == "bullets":
        role = ctx.get("role") or "Role"
        company = ctx.get("company") or "Company"
        bullets = [
            f"Drove end-to-end initiatives as {role} at {company}, improving key KPIs by 15%+.",
            "Collaborated cross-functionally to ship features on time while reducing defects.",
            "Automated repetitive workflows to save team 4–6 hrs/week and increase consistency.",
            "Translated business goals into actionable plans with clear milestones and metrics.",
        ]
        return {"bullets": bullets}

    return JSONResponse({"detail": "Unsupported type"}, status_code=400)


def _format_text_resume(data: ResumeData) -> str:
    parts: List[str] = []
    header = data.name or "Your Name"
    sub = " | ".join([x for x in [data.title, data.email, data.phone, data.location] if x])
    parts.append(header)
    if sub:
        parts.append(sub)
    if data.summary:
        parts.append("\nSummary\n" + data.summary)

    if data.experience:
        parts.append("\nExperience")
        for e in data.experience:
            line = " - ".join([x for x in [e.role, e.company, e.period] if x])
            parts.append(line)
            for b in [b for b in e.bullets if b.strip()]:
                parts.append(f"  • {b}")

    if data.education:
        parts.append("\nEducation")
        for ed in data.education:
            line = " - ".join([x for x in [ed.degree, ed.school, ed.period] if x])
            parts.append(line)
            if ed.details:
                parts.append(f"  • {ed.details}")

    if data.skills:
        skills_line = ", ".join([s.name + (f" ({s.level})" if s.level else "") for s in data.skills if s.name])
        if skills_line:
            parts.append("\nSkills\n" + skills_line)

    if data.achievements:
        parts.append("\nAchievements")
        for a in [a for a in data.achievements if a.strip()]:
            parts.append(f"  • {a}")

    return "\n".join(parts).strip() + "\n"


@app.post("/export/txt")
async def export_txt(payload: ExportPayload):
    text = _format_text_resume(payload.data)
    bio = BytesIO(text.encode("utf-8"))
    headers = {"Content-Disposition": "attachment; filename=resume.txt"}
    return StreamingResponse(bio, media_type="text/plain", headers=headers)


@app.post("/export/docx")
async def export_docx(payload: ExportPayload):
    from docx import Document  # python-docx
    from docx.shared import Pt

    d: ResumeData = payload.data
    doc = Document()

    # Styles
    style = doc.styles['Normal']
    style.font.name = 'Inter'
    style.font.size = Pt(10)

    doc.add_heading(d.name or 'Your Name', level=0)
    sub = " | ".join([x for x in [d.title, d.email, d.phone, d.location] if x])
    if sub:
        p = doc.add_paragraph(sub)
        p.style = doc.styles['Normal']

    if d.summary:
        doc.add_heading('Summary', level=1)
        doc.add_paragraph(d.summary)

    if d.experience:
        doc.add_heading('Experience', level=1)
        for e in d.experience:
            line = " - ".join([x for x in [e.role, e.company, e.period] if x])
            doc.add_paragraph(line)
            for b in [b for b in e.bullets if b.strip()]:
                doc.add_paragraph(b, style=None).style = doc.styles['List Bullet']

    if d.education:
        doc.add_heading('Education', level=1)
        for ed in d.education:
            line = " - ".join([x for x in [ed.degree, ed.school, ed.period] if x])
            doc.add_paragraph(line)
            if ed.details:
                doc.add_paragraph(ed.details)

    if d.skills:
        doc.add_heading('Skills', level=1)
        skills_line = ", ".join([s.name + (f" ({s.level})" if s.level else "") for s in d.skills if s.name])
        doc.add_paragraph(skills_line)

    if d.achievements:
        doc.add_heading('Achievements', level=1)
        for a in [a for a in d.achievements if a.strip()]:
            doc.add_paragraph(a, style=None).style = doc.styles['List Bullet']

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    headers = {"Content-Disposition": "attachment; filename=resume.docx"}
    return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers=headers)


@app.post("/export/pdf")
async def export_pdf(payload: ExportPayload):
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    from reportlab.lib.utils import simpleSplit

    text = _format_text_resume(payload.data)
    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    width, height = letter

    # Simple ATS-friendly PDF
    x_margin = 40
    y = height - 40
    c.setFont("Helvetica-Bold", 14)
    lines = text.split("\n")
    for line in lines:
        if y < 60:
            c.showPage()
            c.setFont("Helvetica", 11)
            y = height - 40
        # Headings heuristic
        if line.strip() in ["Summary", "Experience", "Education", "Skills", "Achievements"]:
            c.setFont("Helvetica-Bold", 12)
            c.drawString(x_margin, y, line)
            c.setFont("Helvetica", 11)
            y -= 18
            continue
        wrapped = simpleSplit(line, "Helvetica", 11, width - 2 * x_margin)
        for w in wrapped:
            c.drawString(x_margin, y, w)
            y -= 14
    c.showPage()
    c.save()
    buf.seek(0)
    headers = {"Content-Disposition": "attachment; filename=resume.pdf"}
    return StreamingResponse(buf, media_type="application/pdf", headers=headers)


@app.get("/test")
def test_database():
    response = {
        "backend": "✅ Running",
        "database": "❌ Not Available",
        "database_url": None,
        "database_name": None,
        "connection_status": "Not Connected",
        "collections": []
    }

    try:
        from database import db  # type: ignore
        if db is not None:
            response["database"] = "✅ Available"
            response["database_url"] = "✅ Configured"
            response["database_name"] = db.name if hasattr(db, 'name') else "✅ Connected"
            response["connection_status"] = "Connected"
            try:
                collections = db.list_collection_names()
                response["collections"] = collections[:10]
                response["database"] = "✅ Connected & Working"
            except Exception as e:
                response["database"] = f"⚠️  Connected but Error: {str(e)[:50]}"
        else:
            response["database"] = "⚠️  Available but not initialized"
    except ImportError:
        response["database"] = "❌ Database module not found (run enable-database first)"
    except Exception as e:
        response["database"] = f"❌ Error: {str(e)[:50]}"

    response["database_url"] = "✅ Set" if os.getenv("DATABASE_URL") else "❌ Not Set"
    response["database_name"] = "✅ Set" if os.getenv("DATABASE_NAME") else "❌ Not Set"

    return response


if __name__ == "__main__":
    import uvicorn
    port = int(os.getenv("PORT", 8000))
    uvicorn.run(app, host="0.0.0.0", port=port)
